#!/usr/bin/env python3
"""Build formatted Word (.docx) journal-submission files from the journal-facing
markdown editions in the `glm writer` folder.

Pipeline per paper:
  1. Parse the `% KEY: value` metadata block.
  2. Strip metadata; run pandoc (markdown -> docx) in the paper folder so that
     figure references resolve; LaTeX math becomes native OMML equations.
  3. Post-process with python-docx:
       - Title page (title, article type, target venue, anonymised author block,
         date, contribution statement, keywords) + single page break.
       - Times New Roman 12 pt body, 1.5 line spacing, justified, 0.5" first-line
         indents; headings bold; captions centred 10 pt; references hanging
         indent; three-line tables at 9.5 pt; running head + page numbers.
       - NO table of contents (formal journal submission).
"""
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Times New Roman"
DATE_STR = "30 August 2026"

META_KEYS = ("TITLE", "VENUE", "TYPE", "RUNNING", "KEYWORDS", "CONTRIBUTION")


def parse_metadata(md_text: str):
    meta, lines, i = {}, md_text.splitlines(), 0
    for ln in lines:
        m = re.match(r"^%\s*([A-Z]+)\s*:\s*(.*)\s*$", ln)
        if m and m.group(1) in META_KEYS:
            meta[m.group(1)] = m.group(2).strip()
            i += 1
        elif ln.strip() == "":
            i += 1
        else:
            break
    body = "\n".join(lines[i:]).strip()
    return meta, body


def add_page_field(paragraph, prefix="Page "):
    run = paragraph.add_run(prefix)
    run.font.name = FONT
    run.font.size = Pt(10)
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(10)
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def set_run_font(run, size=12, bold=None, italic=None, color=None):
    run.font.name = FONT
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def insert_title_page(doc, meta):
    first = doc.paragraphs[0]

    def para(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=6, space_before=0):
        p = first.insert_paragraph_before()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = 1.0
        if text:
            r = p.add_run(text)
            set_run_font(r, size=size, bold=bold, italic=italic)
        return p

    para("", space_after=24)
    para(meta.get("TITLE", ""), size=18, bold=True, space_after=18)
    para(meta.get("TYPE", ""), size=12, italic=True, space_after=6)
    para(f"Prepared for submission to {meta.get('VENUE', '')}", size=12, space_after=30)
    para("", space_after=18)
    para("Author information withheld for double-anonymous peer review.", size=11, italic=True, space_after=4)
    para("Corresponding-author details to be completed by the author at submission.", size=11, italic=True, space_after=30)
    para(DATE_STR, size=12, space_after=24)
    c = meta.get("CONTRIBUTION", "")
    if c:
        p = para("", size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
        r1 = p.add_run("Contribution statement. ")
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(c)
        set_run_font(r2, size=11)
    k = meta.get("KEYWORDS", "")
    if k:
        p = para("", size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)
        r1 = p.add_run("Keywords. ")
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(k)
        set_run_font(r2, size=11)
    # single page break after the cover page
    pb = para("", space_after=0)
    r = pb.add_run()
    r.add_break(WD_BREAK.PAGE)
    return doc


def is_math_only(p):
    """Paragraph consists solely of display math (no visible text)."""
    has_math = p._p.findall(qn("m:oMath")) or p._p.findall(qn("m:oMathPara"))
    if not has_math:
        return False
    return not (p.text or "").strip()


def has_drawing(p):
    return bool(p._p.findall(".//" + qn("w:drawing")))


def is_list_paragraph(p):
    pPr = p._p.pPr
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def caption_kind(text):
    t = text.strip()
    if re.match(r"^(Figure|Table|Fig\.?)\s+\d+", t):
        return "caption"
    return None


def three_line_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    borders = OxmlElement("w:tblBorders")
    for edge, val, sz in (("top", "single", 12), ("bottom", "single", 12),
                          ("left", "none", 0), ("right", "none", 0),
                          ("insideH", "none", 0), ("insideV", "none", 0)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        if val != "none":
            el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:color"), "000000")
        tcB.append(bottom)
        tcPr.append(tcB)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True


def format_document(doc, meta):
    # ---- section-level furniture -------------------------------------
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        # header (not on first/title page): running head, right aligned
        hp = section.header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(meta.get("RUNNING", ""))
        set_run_font(r, size=10)
        # footer: centred page number (not on title page)
        fp = section.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(fp)
        ffp = section.first_page_footer.paragraphs[0]
        ffp.text = ""
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fhp = section.first_page_header.paragraphs[0]
        fhp.text = ""

    # ---- default styles ----------------------------------------------
    styles = doc.styles
    for sname, size, bold in (("Normal", 12, False), ("Body Text", 12, False),
                               ("First Paragraph", 12, False), ("Compact", 12, False)):
        try:
            st = styles[sname]
            st.font.name = FONT
            st.font.size = Pt(size)
            st.font.bold = bold
        except KeyError:
            pass
    for sname, size, bold in (("Heading 1", 14, True), ("Heading 2", 12, True),
                              ("Heading 3", 12, True)):
        try:
            st = styles[sname]
            st.font.name = FONT
            st.font.size = Pt(size)
            st.font.bold = bold
            st.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    # ---- paragraph sweep ----------------------------------------------
    in_references = False
    title_page_len = 0
    # find where body starts: first paragraph after our inserted title page.
    # We detect title-page paragraphs by a marker: they precede the first
    # Heading/Abstract paragraph. Simpler: recompute after insertion below.

    body_paras = doc.paragraphs
    # Identify the first paragraph that begins the pandoc body: it is the first
    # paragraph whose style is not one we inserted. We inserted paragraphs with
    # no style names ('Normal') before the original first paragraph; instead of
    # tracking, we remember insertion count via the page-break marker paragraph.
    # Easiest robust approach: walk from the start; title page is everything up
    # to and including the paragraph containing the page break run.
    pb_index = None
    for i, p in enumerate(body_paras):
        for br in p._p.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                pb_index = i
                break
        if pb_index is not None:
            break

    for idx, p in enumerate(doc.paragraphs):
        if pb_index is not None and idx <= pb_index:
            continue  # title page already formatted
        style = (p.style.name or "").lower()
        text = p.text.strip()

        if style.startswith("heading"):
            level = 1
            if "heading 2" in style or "heading2" in style:
                level = 2
            elif "heading 3" in style or "heading3" in style:
                level = 3
            if text.lower().rstrip(".") == "references" or text.lower() == "references":
                in_references = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.first_line_indent = Inches(0)
            pf.left_indent = Inches(0)
            pf.space_before = Pt(18 if level == 1 else 12)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.5 if level == 1 else 1.5
            for r in p.runs:
                set_run_font(r, size=14 if level == 1 else 12, bold=True)
            continue

        pf = p.paragraph_format
        pf.line_spacing = 1.5

        if has_drawing(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Inches(0)
            pf.space_after = Pt(6)
            for r in p.runs:
                set_run_font(r, size=11)
            continue

        if is_math_only(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Inches(0)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            continue

        if caption_kind(text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Inches(0)
            pf.space_before = Pt(10)
            pf.space_after = Pt(10)
            for r in p.runs:
                set_run_font(r, size=10)
            continue

        if is_list_paragraph(p):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Inches(0)
            pf.space_after = Pt(4)
            for r in p.runs:
                set_run_font(r, size=12)
            continue

        if in_references:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            pf.space_after = Pt(6)
            for r in p.runs:
                set_run_font(r, size=12)
            continue

        # default body paragraph
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0.5)
        pf.space_after = Pt(6)
        for r in p.runs:
            set_run_font(r, size=12)

    # ---- tables --------------------------------------------------------
    for table in doc.tables:
        three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        set_run_font(r, size=9.5)


def build(paper_dir: Path, out_name: str):
    md = paper_dir / "manuscript_journal.md"
    meta, body = parse_metadata(md.read_text(encoding="utf-8"))
    body_file = paper_dir / "_body_tmp.md"
    body_file.write_text(body, encoding="utf-8")
    tmp_docx = paper_dir / "_tmp.docx"
    subprocess.run(
        ["pandoc", str(body_file),
         "-f", "markdown+tex_math_dollars+pipe_tables-yaml_metadata_block",
         "-t", "docx", "-o", str(tmp_docx)],
        check=True, cwd=str(paper_dir),
    )
    doc = Document(str(tmp_docx))
    insert_title_page(doc, meta)
    format_document(doc, meta)
    out = paper_dir / out_name
    doc.save(str(out))
    body_file.unlink()
    tmp_docx.unlink()
    return out, meta


PAPERS = [
    ("paper1_typed_architecture", "Paper1_Typed_Architecture_AcademicPaper_2026-08-30.docx"),
    ("paper2_theorem_atlas", "Paper2_Theorem_Atlas_AcademicPaper_2026-08-30.docx"),
    ("paper3_material_ledgers", "Paper3_Material_Ledgers_AcademicPaper_2026-08-30.docx"),
    ("paper4_delay_dynamics", "Paper4_Delay_Dynamics_AcademicPaper_2026-08-30.docx"),
    ("paper5_sampled_governance", "Paper5_Sampled_Governance_AcademicPaper_2026-08-30.docx"),
    ("wave_e1_cod_forecast_ladder", "WaveE1_Cod_Forecast_Ladder_AcademicPaper_2026-08-30.docx"),
    ("wave_e2_cod_intervention", "WaveE2_Cod_Intervention_AcademicPaper_2026-08-30.docx"),
    ("wave_e3_edwards_forecast_ladder", "WaveE3_Edwards_Forecast_Ladder_AcademicPaper_2026-08-30.docx"),
    ("wave_e4_edwards_intervention", "WaveE4_Edwards_Intervention_AcademicPaper_2026-08-30.docx"),
]


def main():
    base = Path(__file__).resolve().parent
    only = sys.argv[1:] or None
    for folder, out_name in PAPERS:
        if only and folder not in only and out_name not in only:
            continue
        out, meta = build(base / folder, out_name)
        print(f"built: {out.relative_to(base)}  [{meta.get('VENUE','')}]")
    print("done.")


if __name__ == "__main__":
    main()
