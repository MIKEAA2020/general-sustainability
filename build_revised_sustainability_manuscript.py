from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from pathlib import Path

src = Path('revised_sustainability_manuscript.md')
out = Path('revised_sustainability_manuscript.docx')
lines = src.read_text(encoding='utf-8').splitlines()

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)
styles['Normal'].paragraph_format.space_after = Pt(5)
styles['Normal'].paragraph_format.line_spacing = 1.08
for name, size, color in [('Title', 24, '17324D'), ('Subtitle', 14, '3F6680'), ('Heading 1', 16, '17324D'), ('Heading 2', 13, '2F5D73'), ('Heading 3', 11.5, '3F6680')]:
    st = styles[name]
    st.font.name = 'Aptos Display' if name != 'Normal' else 'Aptos'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True

# Footer
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Revised working manuscript · 17 August 2026     ')
run.font.size = Pt(8)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
footer._p.append(fld)

# inline markdown renderer
def add_inline(p, text):
    # Remove markdown links but retain URL in parentheses only if different
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        else:
            p.add_run(part)

# Title page handling
idx = 0
if lines and lines[0].startswith('# '):
    p = doc.add_paragraph(style='Title'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, lines[0][2:])
    idx = 1
    if idx < len(lines) and lines[idx].startswith('## '):
        p = doc.add_paragraph(style='Subtitle'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, lines[idx][3:]); idx += 1
    doc.add_paragraph('')

in_equation = False
in_code = False
buffer = []

def flush_equation(buf):
    if not buf: return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n'.join(buf))
    r.font.name = 'Cambria Math'; r.font.size = Pt(10.5); r.italic = True

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

while idx < len(lines):
    line = lines[idx]
    stripped = line.strip()
    idx += 1

    if stripped == '\\[':
        in_equation = True; buffer = []; continue
    if stripped == '\\]' and in_equation:
        flush_equation(buffer); in_equation = False; buffer = []; continue
    if in_equation:
        buffer.append(stripped); continue

    if stripped.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(buffer)); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
            in_code = False; buffer = []
        else:
            in_code = True; buffer = []
        continue
    if in_code:
        buffer.append(line); continue

    if stripped == '---':
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:color'), 'AABBC5')
        pbdr.append(bottom); pPr.append(pbdr)
        continue
    if not stripped:
        continue

    # Markdown table
    if stripped.startswith('|') and idx < len(lines) and re.match(r'^\|?\s*:?-+', lines[idx].strip()):
        headers = [c.strip() for c in stripped.strip('|').split('|')]
        idx += 1  # separator
        rows=[]
        while idx < len(lines) and lines[idx].strip().startswith('|'):
            rows.append([c.strip() for c in lines[idx].strip().strip('|').split('|')]); idx += 1
        table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
        for j,h in enumerate(headers):
            cell=table.rows[0].cells[j]; set_cell_shading(cell,'DCE8EF'); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cell.paragraphs[0]; add_inline(p,h)
            for r in p.runs: r.bold=True
        for row in rows:
            cells=table.add_row().cells
            for j,val in enumerate(row[:len(headers)]): add_inline(cells[j].paragraphs[0],val)
        doc.add_paragraph('')
        continue

    if line.startswith('# '):
        doc.add_page_break()
        p=doc.add_paragraph(style='Heading 1'); add_inline(p,line[2:]); continue
    if line.startswith('## '):
        p=doc.add_paragraph(style='Heading 1'); add_inline(p,line[3:]); continue
    if line.startswith('### '):
        p=doc.add_paragraph(style='Heading 2'); add_inline(p,line[4:]); continue
    if line.startswith('#### '):
        p=doc.add_paragraph(style='Heading 3'); add_inline(p,line[5:]); continue

    if stripped.startswith('> '):
        p=doc.add_paragraph()
        p.paragraph_format.left_indent=Inches(0.35); p.paragraph_format.right_indent=Inches(0.25)
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
        r=p.add_run(stripped[2:]); r.italic=True; r.font.color.rgb=RGBColor.from_string('2F5D73')
        continue

    m=re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m:
        p=doc.add_paragraph(style='List Number'); add_inline(p,m.group(2)); continue
    if stripped.startswith('- '):
        p=doc.add_paragraph(style='List Bullet'); add_inline(p,stripped[2:]); continue

    p=doc.add_paragraph(); add_inline(p,stripped)

# Core properties
core = doc.core_properties
core.title = 'Toward a General Theory of Sustainability: An Architectural Kernel and Composition Language'
core.subject = 'A typed, transformation-aware architecture for ecological, economic, and social sustainability'
core.author = '[Author name]'
core.comments = 'Revised working manuscript incorporating architectural audits'

doc.save(out)
print(out)
